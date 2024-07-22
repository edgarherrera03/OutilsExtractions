#==============================================================================
# Installation du module Pillow et Python-docx (pour le traitement d'images et du word): NE PAS TOUCHER
#==============================================================================
from .install_packages.check_dependencies import check
API_EXIST = False
try:
    check({
        "bs4": "bs4",
        "PIL": "Pillow",
        "docx": "python-docx"
    })
    from bs4 import BeautifulSoup
    from PIL import Image, ImageEnhance
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    API_EXIST = True
except Exception as e:
    print(f"Erreur dans l'import des modules utilisés : {e}")

#==============================================================================
# Modules à importer : NE PAS TOUCHER
#==============================================================================
from qgis.PyQt.QtCore import QSettings, QTranslator, QCoreApplication
from qgis.PyQt.QtGui import QIcon
import os.path
import os 
from qgis.core import *
from qgis.gui import *
from qgis.PyQt.QtWidgets import *
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QLineEdit, QFileDialog
from io import BytesIO
import requests, zipfile
import sys
import shutil
import numpy as np
import re
#==============================================================================
# Fonctions utilisant des modules non installés dans Qgis par défaut
#==============================================================================
TIMEOUT = 10

def extractionUrlCategories():
    """
    Cette fonction extrait les différents types de données que l'on peut trouver
    sur le site geoservices. Elle a été conçue en fonction de la structure actuelle
    du site web et peut nécessiter des ajustements en cas de modifications.
    """
    url = "https://geoservices.ign.fr/services-web-experts"
    response = requests.get(url)
    extraction_list = {}
    
    if response.status_code != 200:
        print("La requête a échoué. Statut de la réponse :", response.status_code)
        return extraction_list
    if API_EXIST:
        soup = BeautifulSoup(response.content, "html.parser")
        extraction_section = soup.find("div", {"class": "field--items"})
    else: 
        print("Erreur dans l'import du module bs4")
        return 
    # Trouver la section des différentes extractions
    
    
    if not extraction_section:
        print("Section des types d'extraction introuvable sur la page.")
        return extraction_list
    
    extractions = extraction_section.find_all("strong")
    
    if not extractions:
        print("Aucun type d'extraction trouvé sur la page.")
        return extraction_list
    
    for extraction_name in extractions:
        extraction_name_text = extraction_name.text.rstrip(".")
        donnees_list = extractionDonneesCategories(extraction_name_text)
        if donnees_list:
            extraction_list[extraction_name_text] = donnees_list
    
    return extraction_list
def extractionDonneesCategories(nom_categorie):
    """
    Cette fonction permet d'extraire les différentes données disponibles pour
    une catégorie donnée du site geoservices. Elle a été conçue en fonction de
    la structure actuelle du site web et peut nécessiter des ajustements en cas
    de modifications.
    """
    url = "https://geoservices.ign.fr/services-web-experts" + "-" + nom_categorie
    response = requests.get(url)
    donnees_list = {}
    
    if response.status_code == 200:
        if API_EXIST:
            soup = BeautifulSoup(response.content, "html.parser")
            field_items = soup.find_all("div", {"class": "field--item"})
        else:
            print("Erreur dans l'import du module bs4")
            return 
        for field_item in field_items:
            if "WFS" in field_item.text:
                tbody = field_item.find_next("tbody")
                if tbody:
                    tr_items = tbody.find_all("tr")
                    if tr_items:
                        for tr in tr_items:
                            nom_donnees = tr.find("td")
                            if nom_donnees:
                                nom_technique = nom_donnees.find_next("td")
                                donnees_list[nom_donnees.text.strip()] = nom_technique.text.strip()
                        return donnees_list
                    else:
                        print(f"Pas de données disponibles pour {nom_categorie.capitalize()}.")
                        return
                else:
                    print(f"Pas de données disponibles pour {nom_categorie.capitalize()}.")
                    return
        if not donnees_list: 
            print(f"Pas de données WFS disponibles pour {nom_categorie.capitalize()}.")
            return 
    else:
        print(f"La requête a échoué. Impossible d'accéder aux informations {nom_categorie.capitalize()}.")
        return
def ExtractionPlan(extended_clipBbox, dossier, width, height, images_data):
    if API_EXIST:
        api_key = "pk.eyJ1IjoiZWRnYXJoZXJyZXJhMDMiLCJhIjoiY2x3cmpxczNwMDFqajJuczd3aTVxdm1qaiJ9.VIAvZpoMzcnDKduc0P140A"
        base_url = "https://api.mapbox.com/styles/v1/mapbox/streets-v11/static"
        bbox_str = f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}"
        url = f"{base_url}/[{bbox_str}]/{width}x{height}?access_token={api_key}"
        output_file = os.path.join(dossier, "Plan_OSM.png")
        response = requests.get(url)
        if response.status_code == 200:
            image = Image.open(BytesIO(response.content))
            image.save(output_file, "PNG")
            image_data = {
                'title': 'Plan OSM',
                'source': 'OpenStreetMap',
                'path' : output_file,
                'width': 9,
                'height': 9,
                'nom_partie': "Localisation du site",
                'legend': False
            }
            images_data.append(image_data)
        else:
            print(f"La requte a échouée avec le code de statut : {response.status_code}")
            return
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionPlanOrtho(extended_clipBbox, dossier, width, height, images_data):
    if API_EXIST:
        url = 'https://wxs.ign.fr/ortho/inspire/r/wms?'
        params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": "ORTHOIMAGERY.ORTHOPHOTOS",
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Styles": "", 
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
        output_file = os.path.join(dossier, "Plan_Ortho.png")
        response = requests.get(url, params=params)
        if response.status_code == 200:
            image = Image.open(BytesIO(response.content))
            image.save(output_file, "PNG")
            image_data = {
                'title': 'Plan Ortho',
                'source': 'Géoservices',
                'path' : output_file,
                'width': 9,
                'height': 9, 
                'nom_partie': "Localisation du site",
                'legend': False
            }
            images_data.append(image_data)
        else:
            print(f"La requte a échouée avec le code de statut : {response.status_code}")
            return
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionPlanIGN(extended_clipBbox, dossier_resultats, width, height, images_data):
    url = 'https://wxs.ign.fr/cartes/inspire/r/wms?'
    Layer = 'GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2'
    url_legende = "https://data.geopf.fr/annexes/ressources/legendes/" + f"{Layer}" + "-legend.png"    
    params = {
            "Service": "WMS",
            "Request": "GetMap", 
            "Layers": f"{Layer}",
            "Format": "image/png",
            "Version": "1.3.0",
            "CRS": "EPSG:3857",
            "Styles": "GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2", 
            "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
            "Width": f"{width}", 
            "Height": f"{height}"
        }
    dossier_PlanIGN = os.path.join(dossier_resultats, "Plan_IGN.png")
    try:
        response = requests.get(url, params=params, timeout=TIMEOUT)
        if response.status_code == 200:
            image = Image.open(BytesIO(response.content))
            image.save(dossier_PlanIGN, "PNG")
            image_data = {
                'title': 'Plan IGN',
                'source': 'Géoservices',
                'path' : dossier_PlanIGN,
                'width': 9,
                'height': 9,
                'nom_partie': "Localisation du site",
                'legend': True
            }
            images_data.append(image_data)
        else:
            print(f"La requete a échouée avec le code de statut : {response.status_code}")
            return 0
        final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
        layer = QgsRasterLayer(final_url, "Plan IGN")   
        if not layer.isValid():
            print("Erreur lors du chargement de la couche")
        else:
            QgsProject.instance().addMapLayer(layer)
        return dossier_PlanIGN
    except Exception as e:
        print(f"Erreur lors de la récuperation du Plan IGN : {e}")
        return 0
def ExtractionZNIEFF(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'https://ws.carmencarto.fr/WMS/119/fxx_inpn?'
        couches = ["Znieff1","Znieff2"]
        list_img_couches = []
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches: 
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "SRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "", 
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
                
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert("RGBA")
                    
                    # Ajuster la transparence de l'image de couche
                    alpha = image.split()[3]
                    alpha = ImageEnhance.Brightness(alpha).enhance(0.8)  # Réduire l'opacité à 50%
                    image.putalpha(alpha)
                    
                    list_img_couches.append(image)
                else:
                    print(f"La requête a échouée avec le code de statut : {response.status_code}")
                # Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except Exception as e:
                print(f"Pas de données pour ZNIEFF: {e}")

        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.alpha_composite(img)
            
        output_file = os.path.join(dossier, "ZNIEFF.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': 'ZNIEFF',
            'source': 'Inventaire National du Patrimoine Naturel (INPN)',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Espaces remarquables",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionNatura2000(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'https://wxs.ign.fr/environnement/inspire/r/wms?'
        couches = ["PROTECTEDAREAS.SIC","PROTECTEDAREAS.ZPS"]
        list_img_couches = []
        
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "", 
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content))
                    
                    # Ajuster la transparence de l'image de couche
                    alpha = image.split()[3]
                    alpha = ImageEnhance.Brightness(alpha).enhance(0.8)  # Réduire l'opacité à 50%
                    image.putalpha(alpha)
                    
                    list_img_couches.append(image)
                else:
                    print(f"La requte a échouée avec le code de statut : {response.status_code} {couche}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Natura2000: {e}")

        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.alpha_composite(img)
        
        output_file = os.path.join(dossier, "Natura2000.png")
        result_image.save(output_file, "PNG")
        
        image_data = {
            'title': 'Natura 2000',
            'source': 'Géoservices',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Espaces remarquables",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionTopographie(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'https://wxs.ign.fr/topographie/inspire/v/wms?'
        couches = ["BDTOPO-GEOPO-BATI,hedge.hedge,BDTOPO-GEOPO-RESEAU-ROUTIER","BDTOPO-GEOPO-HYDROGRAPHIE,BDTOPO-GEOPO-OROGRAPHIE",
        "TRANSPORTNETWORK.COMMONTRANSPORTELEMENTS.MARKERPOST",
        "BDTOPO-GEOPO-TRANSPORT-ENERGIE,BDTOPO-GEOPO-VEGETATION",
        "BDTOPO-GEOPO-ZONE-ACTIVITES"]
        list_img_couches = []
        
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
                
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content))
                    list_img_couches.append(image)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                    
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Topographie: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
        
        output_file = os.path.join(dossier, "Topographie.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': 'Topographie',
            'source': 'Géoservices',
            'path' : output_file,
            'width': 9,
            'height': 9,
            'nom_partie': "Topographie",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionGeologie(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://geoservices.brgm.fr/geologie?'
        couches = {
            "SCAN_H_RELIEF_GEOL50": "Carte géologique harmonisée",
            "BGM": "Géologie Marine",
            "BSS" : "Banque du sous-sol",
            "LITHO_1M_SIMPLIFIEE": "Carte lithologique simplifiée"
        }
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert("RGBA")
                    image_background = Image.open(dossier_PlanIGN)
                    
                    # Ajuster la transparence de l'image de couche
                    alpha = image.split()[3]
                    alpha = ImageEnhance.Brightness(alpha).enhance(0.8)  # Réduire l'opacité à 50%
                    image.putalpha(alpha)
                    
                    result_image = image_background.convert("RGBA")
                    result_image.alpha_composite(image)
                    output_file = os.path.join(dossier, f"{couches[couche]}.png")
                    result_image.save(output_file, "PNG")
                    image_data = {
                        'title': couches[couche],
                        'source': 'BRGM Info Terre',
                        'path' : output_file,
                        'width': 9,
                        'height': 9, 
                        'nom_partie': "Géologie",
                        'legend': True
                    }
                    images_data.append(image_data)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Geologie: {e}")
    else:
        print("Erreur dans l'import du module Pillow")
        return    
def ExtractionCavites(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://geoservices.brgm.fr/risques?'
        couches = ["CAVITE_LOCALISEE","CAVITE_COMMUNE"]
        list_img_couches = []
    
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert('RGBA')
                    list_img_couches.append(image)
                else:
                    print(f"La requte a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Cavites: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
        
        output_file = os.path.join(dossier, "Cavites.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Cavités",
            'source': 'BRGM Info Terre',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Cavités",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return  
def ExtractionPollution(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://georisques.gouv.fr/services?'
        couches = ["SSP_INSTRUCTION"]
        list_img_couches = []
    
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert('RGBA')
                    list_img_couches.append(image)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Pollution: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
        
        output_file = os.path.join(dossier, "Pollution.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Pollution",
            'source': 'Géorisques',
            'path' : output_file,
            'width': 9,
            'height': 9,
            'nom_partie': "Pollution",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return 
def ExtractionRemonteeNappes(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://georisques.gouv.fr/services?'
        couches = ["REMNAPPE_FR"]
        list_img_couches = []
    
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert('RGBA')
                    list_img_couches.append(image)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Remontée des nappes: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
        
        output_file = os.path.join(dossier, "Remontee_Nappes.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Remontée des nappes",
            'source': 'Géorisques',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Inondation et autre risque naturel",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionInondation(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://georisques.gouv.fr/services?'
        couches = ["PPRN_COMMUNE_RISQINOND_APPROUV", "PPRN_COMMUNE_RISQINOND_PRESCRIT"]
        list_img_couches = []
    
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert('RGBA')
                    list_img_couches.append(image)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Inondations: {e}")

        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
        
        output_file = os.path.join(dossier, "Inondations.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Inondations",
            'source': 'Géorisques',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Inondation et autre risque naturel",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionInondationSumersionMarine(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST: 
        url = 'http://georisques.gouv.fr/services?'
        couches = ["PPRN_COMMUNE_SUBMAR_APPROUV", "PPRN_COMMUNE_SUBMAR_PRESCRIT"]
        list_img_couches = []
    
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert('RGBA')
                    list_img_couches.append(image)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Inondations: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
            
        output_file = os.path.join(dossier, "Inondations_Sumersion_Marine.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Inondations par submersion marine",
            'source': 'Géorisques',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Inondation et autre risque naturel",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionCouleesDeBoue(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://georisques.gouv.fr/services?'
        couches = ["MVT_LOCALISE"]
        list_img_couches = []
        
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}", 
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    image = Image.open(BytesIO(response.content)).convert('RGBA')
                    list_img_couches.append(image)
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Remontée des nappes: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
            
        output_file = os.path.join(dossier, "Coulees_de_boue.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Coulées de boue",
            'source': 'Géorisques',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Inondation et autre risque naturel",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def ExtractionZonesHumides(extended_clipBbox, dossier, width, height, images_data, dossier_PlanIGN):
    if API_EXIST:
        url = 'http://wms.reseau-zones-humides.org/cgi-bin/wmsfma?'
        couches = ["zones_humides_et_plans_d_eau"]
        list_img_couches = []
        
        image_background = Image.open(dossier_PlanIGN)
        
        for couche in couches:
            params = {
                "Service": "WMS",
                "Request": "GetMap", 
                "Layers": couche,
                "Format": "image/png",
                "Version": "1.3.0",
                "CRS": "EPSG:3857",
                "Transparent": "True",
                "Styles": "",
                "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
                "Width": f"{width}",
                "Height": f"{height}"
            }
            try:
                response = requests.get(url, params=params, timeout=TIMEOUT)
                if response.status_code == 200:
                    try:
                        image = Image.open(BytesIO(response.content)).convert('RGBA')
                        list_img_couches.append(image)
                    except Exception as e:
                        print(f"Erreur dans l'extraction de la couche : {e}")
                else:
                    print(f"La requete a échouée avec le code de statut : {response.status_code}")
                #Ajout des couches sur QGIS
                final_url = url + '&'.join([f"{key}={value}" for key, value in params.items()])
                layer = QgsRasterLayer(final_url, f"{couche}")   
                if layer.isValid():
                    QgsProject.instance().addMapLayer(layer)
                else: 
                    print("Erreur lors du chargement de la couche")
            except requests.exceptions.RequestException as e:
                print(f"Pas de données pour Zones humides: {e}")
        
        # Créer une image finale en superposant les couches
        result_image = image_background.convert("RGBA")
        for img in list_img_couches:
            result_image.paste(img, (0, 0), img)
            
        output_file = os.path.join(dossier, "Zones_humides.png")
        result_image.save(output_file, "PNG")
        image_data = {
            'title': "Zones humides",
            'source': 'Inventaire National du Patrimoine Naturel (INPN)',
            'path' : output_file,
            'width': 9,
            'height': 9, 
            'nom_partie': "Espaces remarquables",
            'legend': True
        }
        images_data.append(image_data)
    else:
        print("Erreur dans l'import du module Pillow")
        return
def export_word(document_word, images_data):
    """
    Cette fonction permet de réaliser l'état initial de la faisabilité à partir 
    des images des différents plans qui ont été extraites précédemment. La 
    fonction permet de créer un document Word avec les images correspondantes 
    à chacun des chapitres précisés dans le cahier des charges, accompagnées 
    de leur légende correspondante.
    """
    if not API_EXIST:
        print("Erreur dans l'import d'un des modules.")
        return

    doc = Document()
    parties = []
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Legende")
    images_legend_data = [
        {
            'title':"Plan IGN", 
            'path': os.path.join(path, "PlanIGN_legende.png"),
            "width": 3.99,
            "height": 8.99,
        },
        {
            'title':"Natura 2000", 
            'path': os.path.join(path, "Natura2000_legende.png"),
            "width": 8.74,
            "height": 0.61,
        },
        {
            'title': "Topographie", 
            'path': os.path.join(path, 'topographie_legende.png'),
            "width": 3.99,
            "height": 2.13,
        },
        {
            'title':"Géologie Marine", 
            'path': os.path.join(path, 'BGM_legende.png'),
            "width": 3.99,
            "height": 2.97,
        },
        {
            'title':"Carte géologique harmonisée", 
            'path': os.path.join(path, 'Relief_legende.png'),
            "width": 4.85,
            "height": 0.51,
        },
        {
            'title':"Banque du sous-sol", 
            'path': os.path.join(path, 'Banque_SousSol_legende.png'),
            "width": 3.99,
            "height": 3.66,
        },
        {
            'title':"Carte lithologique simplifiée", 
            'path': os.path.join(path, 'Carte_lithologique_legende.png'),
            "width": 2.49,
            "height": 3.63,
        },
        {
            'title':"Cavités", 
            'path': os.path.join(path, 'Cavites_legende.png'),
            "width": 3.99,
            "height": 3.81,
        },
        {
            'title':"Pollution", 
            'path': os.path.join(path, 'Pollution_legende.png'),
            "width": 8.99,
            "height": 0.64,
        },
        {
            'title':"Remontée des nappes", 
            'path': os.path.join(path, 'Remontee_legende.png'),
            "width": 5.28,
            "height": 9,
        },
        {
            'title':"Inondations", 
            'path': os.path.join(path, 'Inondations_legende.png'),
            "width": 8.99,
            "height": 0.69,
        },
        {
            'title':"Inondations par submersion marine", 
            'path': os.path.join(path, 'Inondations_submersion_legende.png'),
            "width": 13.08,
            "height": 0.69,
        },
        {
            'title':"Coulées de boue", 
            'path': os.path.join(path, 'Coulees_legende.png'),
            "width": 2.82,
            "height": 2.64,
        },
        {
            'title':"Zones humides", 
            'path': os.path.join(path, 'Zones_humides_legende.png'),
            "width": 4.85,
            "height": 2.16,
        },
        {
            'title':"ZNIEFF", 
            'path': os.path.join(path, 'ZNIEFF_legende.png'),
            "width": 8.99,
            "height": 0.81,
        },
    ]
    
    title = doc.add_paragraph(style='Title')
    title.paragraph_format.space_before = Pt(0)
    title_run = title.add_run("État initial - Faisabilité")
    title_run.bold = True  
    title_run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    i = 1

    for img_data in images_data:
        if not img_data['nom_partie'] in parties:
            title_run = doc.add_heading(level=1).add_run(f"{i}. {img_data['nom_partie']}")
            title_run.font.bold = True
            title_run.font.size = Pt(14) 
            parties.append(img_data['nom_partie'])
            i+=1

        title_img = doc.add_paragraph()
        title_img_run = title_img.add_run(img_data['title'])
        title_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_img_run.font.size = Pt(12) 
        title_img_run.font.name = "Arial"
        title_img.paragraph_format.space_after = Pt(5)
        
        image_paragraph = doc.add_paragraph()
        run = image_paragraph.add_run()
        run.add_picture(img_data['path'], width=Cm(img_data['width']), height=Cm(img_data['height']))
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        image_paragraph.paragraph_format.space_after = Pt(0)
        
        if img_data['legend']:
            image_legend_paragraph = doc.add_paragraph()
            run_legend = image_legend_paragraph.add_run()
            supprime_paragraph = True
            for legend_data in images_legend_data:
                if legend_data['title'] == img_data['title']:
                    if legend_data['height'] >= 1:
                        run.add_picture(legend_data['path'], width=Cm(legend_data['width']), height=Cm(legend_data['height']))
                    else:
                        run_legend.add_picture(legend_data['path'], width=Cm(legend_data['width']), height=Cm(legend_data['height']))
                        image_legend_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
                        image_legend_paragraph.paragraph_format.space_after = Pt(0)
                        supprime_paragraph = False
            if supprime_paragraph:
                image_legend_paragraph._element.getparent().remove(image_legend_paragraph._element)

        source_paragraph = doc.add_paragraph()
        source_run = source_paragraph.add_run(f"Source : {img_data['source']}")
        source_run.font.size = Pt(8)
        source_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        source_run.font.color.rgb = RGBColor(51, 102, 153)   

    # Sauvegarder le document Word
    doc.save(document_word)
def ExtractionSimplePlanIGN(extended_clipBbox, dossier, width, height):
    url = 'https://wxs.ign.fr/cartes/inspire/r/wms?'
    Layer = 'GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2'   
    params = {
            "Service": "WMS",
            "Request": "GetMap", 
            "Layers": f"{Layer}",
            "Format": "image/png",
            "Version": "1.3.0",
            "CRS": "EPSG:3857",
            "Styles": "GEOGRAPHICALGRIDSYSTEMS.PLANIGNV2", 
            "BBox": f"{extended_clipBbox[0]},{extended_clipBbox[1]},{extended_clipBbox[2]},{extended_clipBbox[3]}", 
            "Width": f"{width}", 
            "Height": f"{height}"
        }
    dossier_PlanIGN = os.path.join(dossier, "Plan_IGN.png")
    try:
        response = requests.get(url, params=params, timeout=TIMEOUT)
        if response.status_code == 200:
            image = Image.open(BytesIO(response.content))
            image.save(dossier_PlanIGN, "PNG")
        else:
            print(f"La requete a échouée avec le code de statut : {response.status_code}")
            return 0
        return dossier_PlanIGN
    except Exception as e:
        print(f"Erreur lors de la récuperation du Plan IGN : {e}")
        return 0
